"""
modules/report_builder.py
─────────────────────────
Чистый рендерер .docx-отчёта.

Принимает готовый список RegionResult и настройки,
ничего не знает о файловой системе и бизнес-логике.

Публичный API:
    builder = DocxReportBuilder(results, report_date, font="Arial")
    builder.save(Path("/Users/.../Desktop/report.docx"))
"""

import re
import logging
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from modules.result_scanner import RegionResult, ScanStatus

log = logging.getLogger(__name__)


# ══════════════════════════════════════════════════════════════════════════════
#  ПАЛИТРА
# ══════════════════════════════════════════════════════════════════════════════

class Color:
    NAVY      = RGBColor(0x1A, 0x37, 0x5E)
    TEAL      = RGBColor(0x00, 0x7A, 0x87)
    RED       = RGBColor(0xC0, 0x39, 0x2B)
    GRAY_DARK = RGBColor(0x55, 0x66, 0x77)
    GRAY_LITE = RGBColor(0x95, 0xA5, 0xA6)
    WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
    BADGE_SUB = RGBColor(0xAD, 0xC6, 0xE0)   # подписи бейджей
    RED_SOFT  = RGBColor(0xFF, 0x9A, 0x90)   # цифра «ошибок» в шапке


# ══════════════════════════════════════════════════════════════════════════════
#  XML-ХЕЛПЕРЫ  (python-docx не умеет заливку ячеек и бордюры параграфа)
# ══════════════════════════════════════════════════════════════════════════════

def _cell_bg(cell, hex_color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _para_border_bottom(paragraph, color_hex: str = "007A87", size: int = 6) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(size))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color_hex)
    pb.append(bot)
    pPr.append(pb)


# ══════════════════════════════════════════════════════════════════════════════
#  ПОСТРОИТЕЛЬ ДОКУМЕНТА
# ══════════════════════════════════════════════════════════════════════════════

class DocxReportBuilder:
    """Строит .docx-отчёт из списка RegionResult."""

    def __init__(
        self,
        results:     list[RegionResult],
        report_date: str,
        font:        str = "Arial",
    ):
        self.results     = results
        self.report_date = report_date
        self.font        = font
        self.doc         = Document()
        self._setup_page()
        self._setup_styles()

    # ── Инициализация ─────────────────────────────────────────────────────────

    def _setup_page(self) -> None:
        s = self.doc.sections[0]
        s.page_width = s.page_height = None          # сбросим
        s.page_width    = Cm(21)
        s.page_height   = Cm(29.7)
        s.left_margin   = s.right_margin  = Cm(2.5)
        s.top_margin    = s.bottom_margin = Cm(2.0)

    def _setup_styles(self) -> None:
        normal = self.doc.styles["Normal"]
        normal.font.name = self.font
        normal.font.size = Pt(10.5)

    # ── Примитивы ─────────────────────────────────────────────────────────────

    def _run(self, paragraph, text: str, *,
             size: float = 10.5, bold: bool = False, italic: bool = False,
             color: RGBColor = Color.NAVY) -> object:
        r = paragraph.add_run(text)
        r.font.name   = self.font
        r.font.size   = Pt(size)
        r.font.bold   = bold
        r.font.italic = italic
        r.font.color.rgb = color
        return r

    def _new_para(self, *, before: float = 0, after: float = 0,
                  indent: float = 0,
                  align=WD_PARAGRAPH_ALIGNMENT.LEFT) -> object:
        p = self.doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after  = Pt(after)
        if indent:
            p.paragraph_format.left_indent = Cm(indent)
        return p

    def _spacer(self, pt: float = 6) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(pt)

    def _divider(self, color_hex: str = "007A87", size: int = 4) -> None:
        p = self._new_para(before=4, after=8)
        _para_border_bottom(p, color_hex, size)

    # ── Секция: шапка + статистика ────────────────────────────────────────────

    def _build_header(self) -> None:
        """
        Компактная двухколоночная шапка (~1.5 см высотой):
          левая — название и дата
          правая — три inline-бейджа со статистикой
        """
        found  = sum(1 for r in self.results if r.status == ScanStatus.FOUND)
        empty  = sum(1 for r in self.results if r.status == ScanStatus.EMPTY)
        failed = sum(1 for r in self.results if r.status == ScanStatus.FAILED)
        total  = len(self.results)

        table = self.doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Левая: название + подстрочник
        lc = table.cell(0, 0)
        _cell_bg(lc, "1A375E")

        p1 = lc.paragraphs[0]
        p1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p1.paragraph_format.space_before = Pt(7)
        p1.paragraph_format.space_after  = Pt(0)
        p1.paragraph_format.left_indent  = Pt(8)
        self._run(p1, "МОНИТОРИНГ СМИ",
                  size=13, bold=True, color=Color.WHITE)

        p2 = lc.add_paragraph()
        p2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p2.paragraph_format.space_before = Pt(1)
        p2.paragraph_format.space_after  = Pt(7)
        p2.paragraph_format.left_indent  = Pt(8)
        self._run(p2, f"Итоговый отчёт  ·  {self.report_date}  ·  регионов: {total}",
                  size=8, color=Color.BADGE_SUB)

        # Правая: бейджи
        rc = table.cell(0, 1)
        _cell_bg(rc, "1A375E")

        pb = rc.paragraphs[0]
        pb.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        pb.paragraph_format.space_before = Pt(8)
        pb.paragraph_format.space_after  = Pt(8)
        pb.paragraph_format.right_indent = Pt(6)

        badges = [
            (str(found),  "найдено ",  Color.WHITE),
            (str(failed), "ошибок ",    Color.RED_SOFT),
            (str(empty),  "пусто ",  Color.BADGE_SUB),
        ]
        for num, label, num_color in badges:
            self._run(pb, f" {num}",     size=13, bold=True, color=num_color)
            self._run(pb, f" {label} ",  size=8,  color=Color.BADGE_SUB)

        self._spacer(2)

    # ── Секция: найденные упоминания ──────────────────────────────────────────

    def _build_found(self, items: list[RegionResult]) -> None:
        if not items:
            return

        h = self._new_para(before=8, after=2)
        self._run(h, "🔍  ОТВЕТЫ С ПОТЕНЦИАЛЬНЫМИ УПОМИНАНИЯМИ", size=14,
                  bold=True, color=Color.TEAL)
        _para_border_bottom(h, "007A87", size=4)
        self._spacer(4)

        for idx, item in enumerate(items, 1):
            # Заголовок региона
            p_head = self._new_para(before=10, after=1)
            self._run(p_head, f"{idx:02d}.  ", size=11, bold=True, color=Color.TEAL)
            self._run(p_head, item.region,     size=11, bold=True, color=Color.NAVY)

            # Метаданные
            meta_parts = [item.channel]
            if item.posts_count:
                meta_parts.append(f"{item.posts_count} постов")
            if item.scan_date:
                meta_parts.append(item.scan_date)

            p_meta = self._new_para(after=4)
            self._run(p_meta, "  ·  ".join(meta_parts),
                      size=9, italic=True, color=Color.GRAY_DARK)

            # Текст ответа (уже очищен на уровне данных)
            p_ans = self._new_para(after=4, indent=0.5,
                                   align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
            self._run(p_ans, item.answer, size=10.5, color=Color.NAVY)

            # Тонкий разделитель (кроме последнего)
            if idx < len(items):
                p_div = self._new_para(before=6, after=0)
                _para_border_bottom(p_div, "D0DCE8", size=2)

        self._spacer(12)

    # ── Секция: ошибки ────────────────────────────────────────────────────────

    def _build_failed(self, items: list[RegionResult]) -> None:
        if not items:
            return

        h = self._new_para(before=4, after=2)
        self._run(h, "⚠  ОШИБКИ ОБРАБОТКИ — ТРЕБУЕТСЯ ПЕРЕЗАПУСК",
                  size=11, bold=True, color=Color.RED)
        _para_border_bottom(h, "C0392B", size=3)
        self._spacer(4)

        for item in items:
            p = self._new_para(after=2)
            self._run(p, f"▸  {item.region}", size=10, color=Color.RED)
            if item.channel and item.channel != "—":
                self._run(p, f"  ({item.channel})",
                          size=9, italic=True, color=Color.GRAY_DARK)

        self._spacer(12)

    # ── Секция: пустые регионы ────────────────────────────────────────────────

    def _build_empty(self, items: list[RegionResult]) -> None:
        if not items:
            return

        h = self._new_para(before=4, after=2)
        self._run(h, "✓  УПОМИНАНИЙ НЕ НАЙДЕНО",
                  size=11, bold=True, color=Color.GRAY_DARK)
        _para_border_bottom(h, "95A5A6", size=2)
        self._spacer(4)

        p = self._new_para(after=8)
        self._run(p, ", ".join(r.region for r in items),
                  size=9.5, color=Color.GRAY_LITE)

    # ── Футер ─────────────────────────────────────────────────────────────────

    def _build_footer(self) -> None:
        """Футер секции — всегда прижат к низу каждой страницы."""
        section = self.doc.sections[0]
        footer = section.footer

        # Очищаем дефолтный пустой параграф
        for p in footer.paragraphs:
            p.clear()

        p = footer.paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        _para_border_bottom(p, "D0DCE8", size=2)

        ts = datetime.now().strftime("%d.%m.%Y  %H:%M")
        r = p.add_run(f"Отчёт сформирован автоматически  ·  {ts}")
        r.font.name = self.font
        r.font.size = Pt(8.5)
        r.font.italic = True
        r.font.color.rgb = RGBColor(0xB0, 0xB8, 0xC1)

    # ── Публичный API ─────────────────────────────────────────────────────────

    def build(self) -> Document:
        found  = [r for r in self.results if r.status == ScanStatus.FOUND]
        failed = [r for r in self.results if r.status == ScanStatus.FAILED]
        empty  = [r for r in self.results if r.status == ScanStatus.EMPTY]

        self._build_header()
        self._build_found(found)
        self._build_failed(failed)
        self._build_empty(empty)
        self._build_footer()

        return self.doc

    def save(self, path: Path) -> None:
        self.build().save(str(path))
        log.info("Отчёт сохранён → %s", path)
